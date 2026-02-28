"""
Generate Research Methodology section as a DOCX file for the final thesis.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

def set_cell_shading(cell, color_hex):
    """Set background shading for a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shd)

def add_heading_styled(doc, text, level):
    """Add a heading with consistent styling."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_body(doc, text, bold=False, italic=False):
    """Add a body paragraph."""
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_body_continued(para, text, bold=False, italic=False):
    """Add a run to an existing paragraph."""
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    return run

def add_bullet(doc, text, level=0):
    """Add a bullet point."""
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 1.27)
    return p

def create_table(doc, headers, rows, col_widths=None):
    """Create a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
        set_cell_shading(cell, 'D9E2F3')

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def main():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    # =========================================================================
    # TITLE
    # =========================================================================
    add_heading_styled(doc, 'Research Methodology', level=1)

    # =========================================================================
    # INTRODUCTION
    # =========================================================================
    doc.add_paragraph(
        'This chapter describes the research methodology adopted for developing an RT-DETR-based '
        'Explainable Computer-Aided Diagnosis (CAD) system for the automated detection and '
        'classification of peripheral white blood cells (WBCs) from microscopic blood smear images. '
        'The methodology encompasses the research design, dataset description and preparation, '
        'model architecture design, training strategy, evaluation framework, and explainability '
        'analysis. All experiments were conducted using the Ultralytics RT-DETR framework with '
        'PyTorch 2.6.0 on a CUDA-enabled GPU environment.'
    )

    # =========================================================================
    # 1. RESEARCH DESIGN
    # =========================================================================
    add_heading_styled(doc, '1. Research Design', level=2)
    doc.add_paragraph(
        'This study follows a quantitative, experimental research design. The primary objective '
        'is to evaluate the effectiveness of the Real-Time Detection Transformer (RT-DETR) '
        'architecture for WBC detection and classification, with a focus on comparing multiple '
        'backbone architectures to identify the optimal balance between accuracy, inference speed, '
        'and computational efficiency. The research adopts a comparative multi-model evaluation '
        'approach, where four RT-DETR variants with different backbone networks are trained on '
        'the same dataset and evaluated using identical metrics and protocols.'
    )
    doc.add_paragraph(
        'The four RT-DETR model variants investigated in this study are:'
    )
    add_bullet(doc, 'RT-DETR-L: Uses a ResNet-50 backbone with pretrained weights (fine-tuning).')
    add_bullet(doc, 'RT-DETR-X: Uses a ResNet-101 backbone with pretrained weights (fine-tuning).')
    add_bullet(doc, 'RT-DETR-MobileNet: Uses a custom MobileNetV3-Small-inspired backbone, trained from scratch.')
    add_bullet(doc, 'RT-DETR-ShuffleNet: Uses a custom ShuffleNetV2-Small-inspired backbone, trained from scratch.')
    doc.add_paragraph(
        'The first two models (RT-DETR-L and RT-DETR-X) leverage standard pretrained backbones '
        'provided by the Ultralytics framework, while the latter two (RT-DETR-MobileNet and '
        'RT-DETR-ShuffleNet) employ custom lightweight backbone architectures designed specifically '
        'for resource-constrained deployment scenarios such as portable medical devices.'
    )

    # =========================================================================
    # 2. DATASET
    # =========================================================================
    add_heading_styled(doc, '2. Dataset Description', level=2)
    doc.add_paragraph(
        'The Raabin-WBC dataset (Kouzehkanan et al., n.d.) was used for training and evaluation '
        'in this study. Raabin-WBC is a large, freely accessible dataset of white blood cells '
        'collected from normal peripheral blood smears. The dataset was originally compiled using '
        'multiple microscopes and cameras to ensure diversity, and cell annotations were verified '
        'by two independent experts. The dataset contains images of five major WBC types that are '
        'clinically relevant for hematological diagnosis.'
    )

    doc.add_paragraph(
        'The five WBC classes and their corresponding class identifiers used in this study are:'
    )

    # Class mapping table
    create_table(doc,
        headers=['Class Name', 'Class ID', 'Description'],
        rows=[
            ['Basophil', '0', 'Least common WBC; involved in allergic reactions and inflammation'],
            ['Eosinophil', '1', 'Combats parasitic infections and modulates allergic responses'],
            ['Lymphocyte', '2', 'Central to adaptive immunity (T-cells, B-cells, NK cells)'],
            ['Monocyte', '3', 'Largest WBC type; differentiates into macrophages and dendritic cells'],
            ['Neutrophil', '4', 'Most abundant WBC; first responders to bacterial infection'],
        ],
        col_widths=[3.5, 2, 10]
    )
    doc.add_paragraph('')  # spacer

    add_heading_styled(doc, '2.1 Dataset Split', level=3)
    doc.add_paragraph(
        'The dataset was pre-organised into separate training and validation directories, each '
        'containing class-specific subdirectories for both images and their corresponding label files. '
        'The distribution of images across the training and validation sets is presented in the table below.'
    )

    # Dataset distribution table
    create_table(doc,
        headers=['WBC Class', 'Training Images', 'Validation Images', 'Total'],
        rows=[
            ['Basophil', '212', '71', '283'],
            ['Eosinophil', '744', '305', '1,049'],
            ['Lymphocyte', '2,423', '1,017', '3,440'],
            ['Monocyte', '562', '217', '779'],
            ['Neutrophil', '6,236', '2,651', '8,887'],
            ['Total', '10,177', '4,261', '14,438'],
        ],
        col_widths=[3.5, 3.5, 3.5, 3]
    )
    doc.add_paragraph('')

    p = doc.add_paragraph(
        'The dataset exhibits a natural class imbalance that reflects real-world WBC distributions '
        'in peripheral blood. Neutrophils constitute the largest class (61.5%), followed by '
        'Lymphocytes (23.8%), Eosinophils (7.3%), Monocytes (5.4%), and Basophils (2.0%). This '
        'imbalance was intentionally preserved during training to simulate realistic clinical '
        'conditions, consistent with the approach adopted by Katar and Yildirim (2023).'
    )

    add_heading_styled(doc, '2.2 Annotation Format', level=3)
    doc.add_paragraph(
        'Each image in the dataset is accompanied by a corresponding text file containing polygon '
        'segmentation annotations in YOLO format. Each annotation file contains a single line with '
        'the class identifier followed by a series of normalised x, y coordinate pairs defining the '
        'polygon boundary of the cell. During training, the Ultralytics framework automatically '
        'converts these polygon annotations into bounding box coordinates by computing the minimum '
        'enclosing rectangle of the polygon vertices. The YOLO annotation format is as follows:'
    )
    p = doc.add_paragraph()
    run = p.add_run('<class_id> <x1> <y1> <x2> <y2> ... <xn> <yn>')
    run.italic = True
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    doc.add_paragraph(
        'where all coordinate values are normalised to the range [0, 1] relative to the image '
        'dimensions.'
    )

    # =========================================================================
    # 3. DATA PREPROCESSING AND AUGMENTATION
    # =========================================================================
    add_heading_styled(doc, '3. Data Preprocessing and Augmentation', level=2)
    doc.add_paragraph(
        'All images were resized to 512 x 512 pixels during training and inference. The Ultralytics '
        'framework applies a suite of real-time data augmentation techniques during training to '
        'improve model generalisation and mitigate overfitting. The augmentation strategies applied '
        'in this study are summarised below.'
    )

    create_table(doc,
        headers=['Augmentation Technique', 'Parameter', 'Value'],
        rows=[
            ['HSV Hue Shift', 'hsv_h', '0.015'],
            ['HSV Saturation Shift', 'hsv_s', '0.7'],
            ['HSV Value Shift', 'hsv_v', '0.4'],
            ['Translation', 'translate', '0.1'],
            ['Scaling', 'scale', '0.5'],
            ['Horizontal Flip', 'fliplr', '0.5 (probability)'],
            ['Mosaic', 'mosaic', '1.0 (enabled)'],
            ['Random Erasing', 'erasing', '0.4 (probability)'],
            ['Auto Augment', 'auto_augment', 'RandAugment'],
        ],
        col_widths=[5, 3.5, 4]
    )
    doc.add_paragraph('')

    doc.add_paragraph(
        'Mosaic augmentation combines four training images into a single composite image, which '
        'helps the model learn to detect objects at varying scales and in different contextual '
        'arrangements. RandAugment applies a series of randomly selected transformations from a '
        'predefined set of augmentation operations. HSV colour space perturbations account for '
        'staining variations commonly observed in microscopic blood smear images prepared at '
        'different laboratories.'
    )

    # =========================================================================
    # 4. MODEL ARCHITECTURE
    # =========================================================================
    add_heading_styled(doc, '4. Model Architecture', level=2)

    add_heading_styled(doc, '4.1 RT-DETR Overview', level=3)
    doc.add_paragraph(
        'The Real-Time Detection Transformer (RT-DETR), proposed by Zhao et al. (n.d.), is a '
        'transformer-based object detection architecture that achieves real-time performance '
        'while maintaining accuracy comparable to state-of-the-art YOLO models. Unlike YOLO-based '
        'detectors that rely on anchor boxes and Non-Maximum Suppression (NMS) for post-processing, '
        'RT-DETR adopts an end-to-end detection paradigm that eliminates these hand-crafted components.'
    )

    doc.add_paragraph('The RT-DETR architecture comprises four main components:')
    add_bullet(doc,
        'Backbone: Extracts multi-scale feature maps from the input image at different '
        'resolutions (P3, P4, P5 corresponding to 1/8, 1/16, and 1/32 of the original image size).'
    )
    add_bullet(doc,
        'Hybrid Encoder: Consists of two sub-modules \u2014 the Attention-based Intra-scale Feature '
        'Interaction (AIFI) module, which applies self-attention within the highest-level feature map '
        'to capture global context, and the Cross-scale Channel Fusion (CCFF) module, which merges '
        'features from different scales through a top-down and bottom-up Feature Pyramid Network (FPN) '
        'using RepC3 blocks.'
    )
    add_bullet(doc,
        'Uncertainty-Minimal Query Selection: Selects the top-K most informative features as initial '
        'queries for the decoder based on their localisation confidence scores, reducing redundancy '
        'and improving detection quality.'
    )
    add_bullet(doc,
        'Transformer Decoder: Iteratively refines the selected queries through multiple decoder layers '
        'with cross-attention, self-attention, and feed-forward networks to predict the final bounding '
        'boxes and class labels. Position embeddings encode spatial information to enable precise '
        'object localisation.'
    )

    add_heading_styled(doc, '4.2 Pretrained Backbone Variants', level=3)
    doc.add_paragraph(
        'Two standard RT-DETR configurations with pretrained backbone networks were employed:'
    )

    p = doc.add_paragraph()
    add_body_continued(p, 'RT-DETR-L (ResNet-50 backbone): ', bold=True)
    add_body_continued(p,
        'This variant uses a ResNet-50 backbone pretrained on the COCO object detection dataset. '
        'ResNet-50 contains 50 layers with residual skip connections that mitigate the vanishing '
        'gradient problem. The pretrained weights provide robust low-level and mid-level feature '
        'representations, enabling effective transfer learning for the WBC classification task.'
    )

    p = doc.add_paragraph()
    add_body_continued(p, 'RT-DETR-X (ResNet-101 backbone): ', bold=True)
    add_body_continued(p,
        'This variant employs a deeper ResNet-101 backbone, also pretrained on COCO. With 101 layers, '
        'it offers greater representational capacity than ResNet-50, potentially capturing more complex '
        'patterns at the cost of increased computational requirements and memory consumption.'
    )

    add_heading_styled(doc, '4.3 Custom Lightweight Backbone Variants', level=3)
    doc.add_paragraph(
        'To address the objective of developing a lightweight, resource-efficient model suitable '
        'for deployment on portable medical devices, two custom backbone architectures were designed '
        'and integrated with the RT-DETR detection head. These models were trained entirely from '
        'scratch without pretrained weights.'
    )

    p = doc.add_paragraph()
    add_body_continued(p, 'RT-DETR-MobileNet (MobileNetV3-Small-inspired backbone): ', bold=True)
    add_body_continued(p,
        'This custom backbone draws inspiration from the MobileNetV3-Small architecture and employs '
        'depthwise separable convolutions to significantly reduce computational cost. The backbone '
        'consists of a stem convolution layer followed by six stages of depthwise convolution (DWConv) '
        'and pointwise convolution (1x1 Conv) blocks with progressively increasing channel dimensions '
        '(16 \u2192 24 \u2192 40 \u2192 80 \u2192 160). C2f (Cross Stage Partial with two convolutions) '
        'modules are integrated at key stages for enhanced feature aggregation. The detection head uses '
        'an AIFI transformer module with 512 hidden dimensions and 4 attention heads, followed by '
        'RepC3 blocks for multi-scale feature fusion, and an RTDETRDecoder with 300 queries, 4 decoder '
        'layers, 8 attention heads, and 3 feature levels.'
    )

    p = doc.add_paragraph()
    add_body_continued(p, 'RT-DETR-ShuffleNet (ShuffleNetV2-Small-inspired backbone): ', bold=True)
    add_body_continued(p,
        'This custom backbone is inspired by ShuffleNetV2-Small and utilises depthwise separable '
        'convolutions with efficient channel operations. The architecture features a stem convolution '
        'followed by four stages with progressive channel expansion (24 \u2192 48 \u2192 96 \u2192 192 '
        '\u2192 384), using depthwise convolutions for spatial processing and pointwise convolutions '
        'for channel mixing. C2f blocks are used at each stage for feature refinement. The detection '
        'head mirrors the MobileNet variant with identical AIFI, RepC3, and RTDETRDecoder configurations.'
    )

    # Architecture comparison table
    doc.add_paragraph('')
    create_table(doc,
        headers=['Component', 'RT-DETR-L', 'RT-DETR-X', 'RT-DETR-MobileNet', 'RT-DETR-ShuffleNet'],
        rows=[
            ['Backbone', 'ResNet-50', 'ResNet-101', 'MobileNetV3-Small\n(custom)', 'ShuffleNetV2-Small\n(custom)'],
            ['Pretrained', 'Yes (COCO)', 'Yes (COCO)', 'No (from scratch)', 'No (from scratch)'],
            ['Key Operations', 'Residual blocks', 'Residual blocks', 'Depthwise separable\nconvolutions, C2f', 'Depthwise separable\nconvolutions, C2f'],
            ['AIFI Dimensions', '1024, 8 heads', '1024, 8 heads', '512, 4 heads', '512, 4 heads'],
            ['Decoder Queries', '300', '300', '300', '300'],
            ['Decoder Layers', '6', '6', '4', '4'],
            ['Feature Levels', '3 (P3, P4, P5)', '3 (P3, P4, P5)', '3 (P3, P4, P5)', '3 (P3, P4, P5)'],
        ],
        col_widths=[3, 2.8, 2.8, 3.2, 3.2]
    )
    doc.add_paragraph('')

    # =========================================================================
    # 5. TRAINING STRATEGY
    # =========================================================================
    add_heading_styled(doc, '5. Training Strategy', level=2)

    add_heading_styled(doc, '5.1 Training Environment', level=3)
    doc.add_paragraph(
        'All models were trained using the Ultralytics RT-DETR framework (version compatible with '
        'PyTorch 2.6.0 and CUDA 12.4) on a system equipped with a CUDA-enabled GPU with 8 GB VRAM. '
        'Automatic Mixed Precision (AMP) training was enabled for all experiments to reduce memory '
        'consumption and accelerate training without compromising model accuracy. The training pipeline '
        'was implemented in Python 3.12 using Jupyter Notebook as the development environment.'
    )

    add_heading_styled(doc, '5.2 Training Hyperparameters', level=3)
    doc.add_paragraph(
        'Two distinct hyperparameter configurations were used, one for the pretrained models '
        '(RT-DETR-L and RT-DETR-X) and another for the from-scratch models (RT-DETR-MobileNet '
        'and RT-DETR-ShuffleNet). The pretrained models required fewer epochs and could tolerate '
        'a higher initial learning rate due to their pre-learned feature representations. In contrast, '
        'the from-scratch models needed more training epochs, a lower initial learning rate, and '
        'longer warmup periods for stable convergence.'
    )

    create_table(doc,
        headers=['Hyperparameter', 'Pretrained Models\n(RT-DETR-L, RT-DETR-X)', 'From-Scratch Models\n(MobileNet, ShuffleNet)'],
        rows=[
            ['Total Epochs', '5', '15'],
            ['Image Size', '512 x 512', '512 x 512'],
            ['Batch Size', '8 (L), 3 (X)', '12'],
            ['Initial Learning Rate (lr0)', '0.01', '0.0001'],
            ['Final Learning Rate (lrf)', '0.0001', '0.01'],
            ['Momentum', '0.937', '0.937'],
            ['Weight Decay', '0.0005', '0.0005'],
            ['Warmup Epochs', '1', '3'],
            ['Warmup Momentum', '0.8', '0.8'],
            ['Warmup Bias LR', '0.1', '0.01'],
            ['Learning Rate Scheduler', 'Cosine Annealing', 'Cosine Annealing'],
            ['Early Stopping Patience', '15 epochs', '20 epochs'],
            ['Mixed Precision (AMP)', 'Enabled', 'Enabled'],
            ['Workers', '4', '4'],
        ],
        col_widths=[4.5, 4.5, 4.5]
    )
    doc.add_paragraph('')

    doc.add_paragraph(
        'The cosine annealing learning rate scheduler was used for all models to provide a smooth, '
        'gradual decay of the learning rate from the initial value to the final value over the '
        'course of training. This approach helps the model converge to a better optimum compared '
        'to step-based or linear decay schedules.'
    )

    add_heading_styled(doc, '5.3 Loss Function', level=3)
    doc.add_paragraph(
        'The RT-DETR training objective is a weighted combination of three loss components:'
    )
    add_bullet(doc,
        'Box Loss (weight = 7.5): Measures the regression error between predicted and '
        'ground-truth bounding boxes using a combination of L1 loss and Generalised IoU (GIoU) loss.'
    )
    add_bullet(doc,
        'Classification Loss (weight = 0.5): A focal loss variant that measures the classification '
        'error for assigning the correct WBC class to each detected object.'
    )
    add_bullet(doc,
        'Distribution Focal Loss (DFL, weight = 1.5): Refines the bounding box regression by '
        'learning a distribution over possible box coordinates rather than point estimates.'
    )

    add_heading_styled(doc, '5.4 Checkpoint and Resume Training', level=3)
    doc.add_paragraph(
        'A checkpoint management system was implemented to support incremental training across '
        'multiple sessions. After each training session, the model weights (last.pt) and training '
        'metadata (including total epochs completed, session history, and timestamps) were saved to '
        'a persistent checkpoint directory. When resuming training, the system automatically loads '
        'the checkpoint weights, disables warmup epochs (since the model is already warmed up), and '
        'continues training from the last completed epoch. This approach was essential given the '
        'computational constraints of the 8 GB VRAM GPU, as it allowed training to be conducted in '
        'manageable sessions of 2\u20135 epochs each.'
    )

    # =========================================================================
    # 6. EVALUATION FRAMEWORK
    # =========================================================================
    add_heading_styled(doc, '6. Evaluation Framework', level=2)

    add_heading_styled(doc, '6.1 Evaluation Protocol', level=3)
    doc.add_paragraph(
        'Each trained model was evaluated on 500 images sampled from the training set '
        '(100 images per class, randomly selected with a fixed seed of 123 for reproducibility). '
        'A confidence threshold of 0.1 was applied during inference. For each image, the model '
        'prediction with the highest confidence score was selected as the final classification. '
        'Images where the model produced no detections were recorded separately.'
    )

    add_heading_styled(doc, '6.2 Evaluation Metrics', level=3)
    doc.add_paragraph(
        'The following metrics were used to comprehensively evaluate model performance:'
    )

    p = doc.add_paragraph()
    add_body_continued(p, 'Accuracy: ', bold=True)
    add_body_continued(p,
        'The proportion of correctly classified samples out of the total valid predictions. '
        'Calculated as: Accuracy = (TP + TN) / (TP + TN + FP + FN), where TP, TN, FP, and FN '
        'denote true positives, true negatives, false positives, and false negatives respectively.'
    )

    p = doc.add_paragraph()
    add_body_continued(p, 'Precision: ', bold=True)
    add_body_continued(p,
        'The proportion of true positive predictions among all positive predictions for a given class. '
        'Precision = TP / (TP + FP). High precision indicates a low false positive rate.'
    )

    p = doc.add_paragraph()
    add_body_continued(p, 'Recall (Sensitivity): ', bold=True)
    add_body_continued(p,
        'The proportion of actual positive samples correctly identified by the model. '
        'Recall = TP / (TP + FN). High recall indicates a low false negative rate.'
    )

    p = doc.add_paragraph()
    add_body_continued(p, 'F1-Score: ', bold=True)
    add_body_continued(p,
        'The harmonic mean of precision and recall, providing a balanced measure of classification '
        'performance. F1 = 2 \u00d7 (Precision \u00d7 Recall) / (Precision + Recall).'
    )

    p = doc.add_paragraph()
    add_body_continued(p, 'Confusion Matrix: ', bold=True)
    add_body_continued(p,
        'A tabular representation showing the distribution of predicted classes against ground-truth '
        'classes. This provides detailed insight into which classes are most commonly confused.'
    )

    p = doc.add_paragraph()
    add_body_continued(p, 'Average Inference Time: ', bold=True)
    add_body_continued(p,
        'The mean time (in milliseconds) required for the model to process a single image and '
        'produce predictions. This metric is critical for assessing real-time deployment feasibility.'
    )

    add_heading_styled(doc, '6.3 Comparative Analysis', level=3)
    doc.add_paragraph(
        'All four models were compared across accuracy, per-class F1-scores, inference time, and '
        'training time using bar charts, confusion matrices, and summary tables generated in the '
        'comparison notebook (Compare_and_Visualize_Results.ipynb). This comparative analysis enables '
        'identification of the optimal model configuration for different deployment scenarios \u2014 '
        'high-accuracy clinical use versus resource-constrained portable deployment.'
    )

    # =========================================================================
    # 7. EXPLAINABILITY ANALYSIS
    # =========================================================================
    add_heading_styled(doc, '7. Explainability Analysis', level=2)
    doc.add_paragraph(
        'To ensure transparency and build clinical trust in the model\u2019s predictions, two '
        'Class Activation Mapping (CAM) techniques are employed to generate visual explanations '
        'of the model\u2019s decision-making process.'
    )

    p = doc.add_paragraph()
    add_body_continued(p, 'Gradient-weighted Class Activation Mapping (GradCAM): ', bold=True)
    add_body_continued(p,
        'GradCAM (Selvaraju et al., 2017) computes the gradients of the target class score with '
        'respect to the feature maps of a selected convolutional layer. These gradients are '
        'globally average-pooled to obtain importance weights for each feature map channel. '
        'The weighted combination of feature maps is passed through a ReLU activation to produce '
        'a coarse heatmap highlighting the regions that most influenced the prediction. GradCAM '
        'provides class-discriminative localisation without requiring architectural modifications.'
    )

    p = doc.add_paragraph()
    add_body_continued(p, 'Score-weighted Class Activation Mapping (ScoreCAM): ', bold=True)
    add_body_continued(p,
        'ScoreCAM (Wang et al., 2020) is a gradient-free alternative that avoids potential issues '
        'with noisy gradients. Instead of using gradients, ScoreCAM uses each activation map as a '
        'mask on the input image, feeds the masked images through the network, and uses the '
        'resulting class confidence scores as weights for the activation maps. This produces '
        'smoother, more reliable saliency maps that clearly highlight the morphological features '
        '(such as nucleus shape, cytoplasm characteristics, and granularity) that the model relies '
        'upon for classification.'
    )

    doc.add_paragraph(
        'These explainability visualisations serve a dual purpose: (1) they enable hematologists '
        'to verify that the model focuses on clinically relevant cell features (e.g., nucleus '
        'morphology and cytoplasmic characteristics) consistent with expert diagnostic criteria, '
        'and (2) they provide interpretable evidence that supports the model\u2019s predictions, '
        'thereby facilitating clinical adoption and trust.'
    )

    # =========================================================================
    # 8. SOFTWARE AND TOOLS
    # =========================================================================
    add_heading_styled(doc, '8. Software and Tools', level=2)

    create_table(doc,
        headers=['Tool / Library', 'Version', 'Purpose'],
        rows=[
            ['Python', '3.12', 'Programming language'],
            ['PyTorch', '2.6.0 (CUDA 12.4)', 'Deep learning framework'],
            ['Ultralytics', 'Latest', 'RT-DETR model training and inference'],
            ['timm', 'Latest', 'Pretrained model components'],
            ['scikit-learn', 'Latest', 'Evaluation metrics (accuracy, confusion matrix, classification report)'],
            ['OpenCV', 'Latest', 'Image loading and preprocessing'],
            ['Matplotlib / Seaborn', 'Latest', 'Visualisation of results, confusion matrices, and bar charts'],
            ['Jupyter Notebook', '\u2014', 'Interactive development environment'],
            ['Grad-CAM / Score-CAM', '\u2014', 'Explainability heatmap generation'],
        ],
        col_widths=[4, 3, 7]
    )
    doc.add_paragraph('')

    # =========================================================================
    # 9. METHODOLOGY FLOWCHART (textual)
    # =========================================================================
    add_heading_styled(doc, '9. Methodology Summary', level=2)
    doc.add_paragraph(
        'The overall research methodology can be summarised in the following sequential steps:'
    )
    doc.add_paragraph('Step 1: Dataset Preparation', style='List Number')
    add_bullet(doc, 'Organised Raabin-WBC dataset into training and validation splits with class-specific subdirectories.')
    add_bullet(doc, 'Verified and corrected polygon annotation labels to ensure consistency with class identifiers.')
    add_bullet(doc, 'Generated YAML configuration files for the Ultralytics training pipeline.')

    doc.add_paragraph('Step 2: Model Architecture Design', style='List Number')
    add_bullet(doc, 'Configured two standard RT-DETR variants (L and X) with pretrained ResNet backbones.')
    add_bullet(doc, 'Designed two custom lightweight backbone architectures (MobileNetV3-Small and ShuffleNetV2-Small inspired) with depthwise separable convolutions and C2f modules.')
    add_bullet(doc, 'Integrated all backbones with the RT-DETR hybrid encoder and transformer decoder.')

    doc.add_paragraph('Step 3: Model Training', style='List Number')
    add_bullet(doc, 'Trained pretrained models (RT-DETR-L and RT-DETR-X) for 5 total epochs with higher learning rates.')
    add_bullet(doc, 'Trained from-scratch models (RT-DETR-MobileNet and RT-DETR-ShuffleNet) for 15 total epochs with lower learning rates and extended warmup.')
    add_bullet(doc, 'Applied data augmentation (mosaic, HSV shifts, flipping, random erasing, RandAugment) during training.')
    add_bullet(doc, 'Used checkpoint management for incremental training across sessions.')

    doc.add_paragraph('Step 4: Model Evaluation', style='List Number')
    add_bullet(doc, 'Evaluated all four models on 500 images (100 per class) using accuracy, precision, recall, F1-score, and inference time.')
    add_bullet(doc, 'Generated confusion matrices and per-class classification reports.')
    add_bullet(doc, 'Compared model performance using bar charts and summary tables.')

    doc.add_paragraph('Step 5: Explainability Analysis', style='List Number')
    add_bullet(doc, 'Applied GradCAM and ScoreCAM to generate visual heatmaps overlaid on input images.')
    add_bullet(doc, 'Analysed whether the model focuses on clinically relevant morphological features.')
    add_bullet(doc, 'Validated explainability outputs against established hematological diagnostic criteria.')

    # =========================================================================
    # SAVE
    # =========================================================================
    output_path = os.path.join(
        r'C:\D drive\mydata\MSML\GitHub\RT-DETR-Based-Explainable-CAD-System-for-Automated-Detection-and-Classification-of-White-Blood-Cells',
        'Research_Methodology.docx'
    )
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path


if __name__ == '__main__':
    main()
